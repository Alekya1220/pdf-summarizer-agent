import streamlit as st
import pdfplumber
import openai
from io import BytesIO
from docx import Document
import speech_recognition as sr
from gtts import gTTS
import tempfile
import os
import time
from langdetect import detect
import pytesseract
from PIL import Image

# -----------------------------
# Streamlit Config
# -----------------------------
st.set_page_config(page_title="PDF Voice Assistant", page_icon="🌍", layout="wide")

# -----------------------------
# Animated Background
# -----------------------------
page_bg = """
<style>
.stApp {
    background: linear-gradient(-45deg, #1a1a2e, #16213e, #0f3460, #533483);
    background-size: 400% 400%;
    animation: gradient 15s ease infinite;
    color: white;
}
@keyframes gradient {
    0% {background-position: 0% 50%;}
    50% {background-position: 100% 50%;}
    100% {background-position: 0% 50%;}
}
.summary-box {
    background: rgba(0,0,0,0.6);
    padding: 20px;
    border-radius: 12px;
    box-shadow: 0px 4px 15px rgba(0,0,0,0.5);
}
</style>
"""
st.markdown(page_bg, unsafe_allow_html=True)

# -----------------------------
# OpenAI API Key
# -----------------------------
openai.api_key = "YOUR_OPENAI_API_KEY"  # Replace with your key

# -----------------------------
# Language Map
# -----------------------------
lang_map = {
    "en": ("en", "en-IN"),
    "te": ("te", "te-IN"),
    "hi": ("hi", "hi-IN"),
    "ta": ("ta", "ta-IN"),
    "kn": ("kn", "kn-IN"),
    "ml": ("ml", "ml-IN"),
    "bn": ("bn", "bn-IN"),
    "gu": ("gu", "gu-IN"),
    "mr": ("mr", "mr-IN"),
    "pa": ("pa", "pa-IN"),
    "ur": ("ur", "ur-IN"),
    "or": ("or", "or-IN"),
}

def detect_language(text):
    try:
        lang_code = detect(text)
        if lang_code in lang_map:
            return lang_map[lang_code]
        else:
            return ("en", "en-IN")
    except:
        return ("en", "en-IN")

# -----------------------------
# Functions
# -----------------------------
def extract_text_pagewise(uploaded_file):
    """Extract text page by page with OCR fallback"""
    pages_text = []
    with pdfplumber.open(uploaded_file) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text()
            if not page_text or page_text.strip() == "":
                try:
                    pil_image = page.to_image(resolution=300).original
                    page_text = pytesseract.image_to_string(
                        pil_image,
                        lang="eng+tel+hin+tam+kan+mal+ben+guj+mar+pan+urd+ori"
                    )
                except Exception as e:
                    page_text = f"[OCR failed on page {i}] {str(e)}"
            pages_text.append(page_text.strip())
    return pages_text

def summarize_section(text, section_name="Section", language="en"):
    prompt = f"""
    Summarize the following text into a bullet-point list.

    Rules:
    - Reply in **{language}**.
    - Keep points short and clear.
    - Ignore unnecessary details.
    - Focus only on main ideas.

    {section_name}:
    {text}
    """
    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": f"You are a helpful assistant that summarizes documents in {language}."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.4,
    )
    return response.choices[0].message["content"]

def answer_question(question, context_text, chat_history, language="en"):
    messages = [{"role": "system", "content": f"You are a helpful assistant that answers only based on the PDF content. Reply in {language}."}]
    messages.extend(chat_history)
    messages.append({"role": "user", "content": f"Context:\n{context_text}\n\nQuestion: {question}"})

    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=messages,
        temperature=0.3,
    )
    return response.choices[0].message["content"]

def export_txt(summary):
    return BytesIO(summary.encode("utf-8"))

def export_docx(summary):
    doc = Document()
    doc.add_heading("PDF Summary", level=1)
    for line in summary.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def speech_to_text(language="en-IN"):
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        st.info("🎙 Listening... Say 'stop' to end voice mode.")
        audio = recognizer.listen(source, phrase_time_limit=8)
    try:
        query = recognizer.recognize_google(audio, language=language)
        return query
    except sr.UnknownValueError:
        return ""
    except sr.RequestError:
        return "Speech recognition unavailable."

def text_to_speech(text, language="en"):
    tts = gTTS(text=text, lang=language)
    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
    tts.save(tmp_file.name)
    return tmp_file.name

# -----------------------------
# Sidebar Options
# -----------------------------
st.sidebar.title("⚙️ Settings")
summary_mode = st.sidebar.radio("Choose Summary Mode:", ("📑 Page-wise Summary", "📘 Full Document Summary"))
qa_mode = st.sidebar.checkbox("🤖 Enable Conversational Q&A")
voice_mode = st.sidebar.checkbox("🎤 Enable Continuous Voice Mode")
lang_setting = st.sidebar.radio("🌍 Language Mode:", ("Automatic Detection", "Manual Selection"))

manual_lang = None
if lang_setting == "Manual Selection":
    manual_lang_choice = st.sidebar.selectbox("Choose Language", list(lang_map.keys()))
    manual_lang = lang_map[manual_lang_choice]

# -----------------------------
# Main App
# -----------------------------
st.title("📄 PDF Summarizer + Voice Q&A (Multi-Language + OCR)")
st.write("Upload a PDF (scanned or text), chat with text/voice, and listen to answers in any Indian language 🌍")

uploaded_file = st.file_uploader("Upload a PDF file", type=["pdf"])

if uploaded_file:
    with st.spinner("📖 Extracting text..."):
        pages = extract_text_pagewise(uploaded_file)

    if not any(pages):
        st.error("⚠️ Could not extract text from the PDF.")
    else:
        summaries = []
        if summary_mode == "📑 Page-wise Summary":
            for i, page_text in enumerate(pages, start=1):
                if page_text.strip():
                    with st.spinner(f"Summarizing Page {i}..."):
                        if manual_lang:
                            lang_code, _ = manual_lang
                        else:
                            lang_code = "en"
                        page_summary = summarize_section(page_text, section_name=f"Page {i}", language=lang_code)
                        summaries.append(f"### 📑 Page {i} Summary\n{page_summary}\n")
            final_summary = "\n".join(summaries)
        else:
            full_text = "\n".join(pages)
            with st.spinner("Summarizing full document..."):
                if manual_lang:
                    lang_code, _ = manual_lang
                else:
                    lang_code = "en"
                final_summary = summarize_section(full_text, section_name="Full Document", language=lang_code)

        # Show Summary
        st.markdown('<div class="summary-box">', unsafe_allow_html=True)
        st.subheader("📌 Summary")
        st.markdown(final_summary, unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        # Download
        st.download_button("📥 Download as TXT", data=export_txt(final_summary), file_name="summary.txt")
        st.download_button("📥 Download as DOCX", data=export_docx(final_summary), file_name="summary.docx")

        # Q&A Section
        if qa_mode:
            st.subheader("🤖 Chat with your PDF")

            if "chat_history" not in st.session_state:
                st.session_state.chat_history = []

            if voice_mode:
                if st.button("🎤 Start Continuous Voice Conversation"):
                    context = "\n".join(pages)
                    while True:
                        _, stt_lang = manual_lang if manual_lang else ("en", "en-IN")
                        query = speech_to_text(language=stt_lang)
                        if query.lower() == "stop":
                            st.warning("🛑 Voice mode stopped.")
                            break
                        if query.strip() != "":
                            st.chat_message("user").markdown(query)

                            if manual_lang:
                                tts_lang, _ = manual_lang
                            else:
                                tts_lang, stt_lang = detect_language(query)

                            answer = answer_question(query, context, st.session_state.chat_history, language=tts_lang)

                            st.session_state.chat_history.append({"role": "user", "content": query})
                            st.session_state.chat_history.append({"role": "assistant", "content": answer})

                            st.chat_message("assistant").markdown(answer)

                            audio_file = text_to_speech(answer, language=tts_lang)
                            st.audio(audio_file, format="audio/mp3")
                            os.remove(audio_file)
                        time.sleep(1)

            else:
                user_question = st.chat_input("Ask a question...")
                if user_question:
                    context = "\n".join(pages)

                    if manual_lang:
                        tts_lang, _ = manual_lang
                    else:
                        tts_lang, _ = detect_language(user_question)

                    with st.spinner("Thinking..."):
                        answer = answer_question(user_question, context, st.session_state.chat_history, language=tts_lang)

                    st.session_state.chat_history.append({"role": "user", "content": user_question})
                    st.session_state.chat_history.append({"role": "assistant", "content": answer})

                    st.chat_message("assistant").markdown(answer)

                    audio_file = text_to_speech(answer, language=tts_lang)
                    st.audio(audio_file, format="audio/mp3")
                    os.remove(audio_file)

            # Show History
            for msg in st.session_state.chat_history:
                if msg["role"] == "user":
                    st.chat_message("user").markdown(msg["content"])
                else:
                    st.chat_message("assistant").markdown(msg["content"])
